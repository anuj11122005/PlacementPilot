import os
from docx import Document
from pathlib import Path

def create_resume():
    doc = Document()
    
    # General Info
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("janesmith@example.com | (555) 123-4567 | San Francisco, CA")
    
    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("Highly skilled Senior Python Developer with 8 years of experience building scalable backend services and microservices. Proven ability to lead teams, architect complex systems, and deliver robust software solutions. Passionate about writing clean, maintainable code and mentoring junior developers.")
    
    # Skills
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("Programming Languages: Python, JavaScript, SQL, Go\n"
                      "Frameworks & Libraries: FastAPI, Django, Flask, React, pandas, SQLAlchemy\n"
                      "Cloud & DevOps: AWS (EC2, S3, RDS), Docker, Kubernetes, CI/CD (GitHub Actions, GitLab CI)\n"
                      "Databases: PostgreSQL, MongoDB, Redis")
    
    # Experience
    doc.add_heading("Work Experience", level=1)
    doc.add_heading("Senior Backend Engineer - TechNova Inc.", level=2)
    doc.add_paragraph("January 2019 - Present")
    doc.add_paragraph("- Designed and implemented a high-throughput microservices architecture using FastAPI and PostgreSQL, handling over 10 million requests per day.\n"
                      "- Reduced API latency by 40% through aggressive caching strategies using Redis and optimizing complex SQL queries.\n"
                      "- Led a team of 4 backend engineers, managing sprints, code reviews, and architectural planning.\n"
                      "- Migrated legacy monolithic applications to Docker containers deployed on AWS Kubernetes (EKS).")
    
    doc.add_heading("Software Engineer - WebSolutions LLC", level=2)
    doc.add_paragraph("June 2015 - December 2018")
    doc.add_paragraph("- Developed and maintained multiple RESTful APIs using Django Rest Framework for a variety of client applications.\n"
                      "- Integrated third-party APIs (Stripe, Twilio, SendGrid) to enable payment processing and notifications.\n"
                      "- Wrote comprehensive unit and integration tests using pytest, achieving 95% test coverage.\n"
                      "- Collaborated closely with the frontend team to ensure seamless API integration.")
    
    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Bachelor of Science in Computer Science")
    doc.add_paragraph("University of Technology, Graduated: May 2015")
    
    # Save the document
    fixture_dir = Path(__file__).parent.parent / "tests" / "fixtures"
    fixture_dir.mkdir(parents=True, exist_ok=True)
    out_path = fixture_dir / "realistic_resume.docx"
    doc.save(str(out_path))
    print(f"Created realistic resume at {out_path}")

if __name__ == "__main__":
    create_resume()
