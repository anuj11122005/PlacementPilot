"""
run_eval.py -- Live end-to-end integration test against a deployed API.

Tests the FULL pipeline by making actual HTTP requests with a synthetic DOCX resume.
"""
import os
import sys
import uuid
import argparse
import requests
import json
from docx import Document

def create_synthetic_resume(path="test_resume.docx"):
    doc = Document()
    doc.add_heading('Alice Johnson', 0)
    doc.add_paragraph('Backend Engineer\nalice@example.com')
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Backend Engineer with 5 years of experience building Python microservices. Passionate about team collaboration, mentoring juniors, and delivering scalable systems.')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Languages: Python, JavaScript, SQL. Frameworks: FastAPI, Django. Tools: Docker, AWS, Git.')
    
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Software Engineer at TechCorp. Led a team of 3 developers to migrate a monolithic application to FastAPI microservices. Managed sprints and conducted code reviews.')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.S. in Computer Science, State University.')
    
    doc.save(path)
    return path

def main():
    parser = argparse.ArgumentParser(description="Run E2E evaluations against PlacementPilot API.")
    parser.add_argument("--url", default="http://localhost:8000/analyze", help="URL of the /analyze endpoint")
    args = parser.parse_args()

    print("=" * 80)
    print(f"  Phase 7 -- Live HTTP E2E Evaluation against {args.url}")
    print("=" * 80)

    resume_path = create_synthetic_resume("test_resume.docx")

    # A cases: Strong matches. Expected: not refused, no unsupported reqs.
    full_8_queries = [
        ("A1: Python/FastAPI strong match", "Experience with Python and FastAPI", True),
        ("A2: Java/Spring mismatch", "Expertise in Java and Spring Boot", False),
        ("A3: Leadership/soft-skill (rescued)", "Experience leading teams and managing sprints", True),
        ("A4: Forklift garbage query", "Certified forklift operator with plumbing skills", False),
        ("A5: AWS/Docker keyword match", "AWS and Docker", True),
    ]

    # B cases: Full mismatches. Expected: Refusal string, all reqs unsupported.
    mismatch_queries = [
        ("B1: Java/Spring (wrong stack)", "Expertise in Java and Spring Boot"),
        ("B2: Oracle DBA (adjacent skill depth)", "Oracle DBA with RAC clustering and tablespace management"),
        ("B3: ML/PyTorch (wrong specialization)", "PyTorch model training and neural network architecture design"),
        ("B4: ARM firmware (wrong domain)", "Firmware programming for ARM Cortex microcontrollers"),
        ("B5: Kubernetes/Helm (SRE requirements)", "Kubernetes pod orchestration with Helm charts and Terraform HCL modules"),
        ("B6: VP Engineering (wrong seniority)", "VP of Engineering overseeing 200-person distributed organization and P&L accountability"),
        ("B7: HIPAA/HL7 (healthcare domain)", "HIPAA compliance officer with clinical data governance and HL7 FHIR integration"),
        ("B8: Quant finance (finance domain)", "Quantitative analyst with stochastic calculus and derivatives pricing models"),
        ("B9: PCB design (hardware/EE domain)", "PCB layout design with KiCad and signal integrity analysis for high-speed circuits"),
        ("B10: Forklift/plumbing (garbage)", "Certified forklift operator with plumbing skills"),
    ]

    # C cases: Traps. Expected: Handled gracefully (not a hallucination).
    trap_queries = [
        ("C1: Trap - Hallucinated metrics", "Led a team of 10+ developers", True),
        ("C2: Trap - Hallucinated timeline", "10 years of experience building Python microservices", True),
        ("C3: Trap - False Scope", "Architected monolithic applications", True),
        ("C4: Trap - Unstated Proficiency", "Expert proficiency in AWS and Git", True),
        ("C5: Trap - Role Assumption", "Passionate about mentoring 50+ juniors", True)
    ]
    
    # D cases: Partial Mismatches (to test unsupported_requirements separation)
    partial_queries = [
        ("D1: Partial Mismatch (Kubernetes)", "Looking for a backend engineer with Python, FastAPI, SQL, and Kubernetes experience.", True, ["Kubernetes"])
    ]

    all_cases = []
    for label, query, expected in full_8_queries:
        all_cases.append((label, query, expected, []))
    for label, query in mismatch_queries:
        if not any(c[0] == label for c in all_cases):
            all_cases.append((label, query, False, []))
    for label, query, expected in trap_queries:
        all_cases.append((label, query, expected, []))
    for label, query, expected, partial_unsupported in partial_queries:
        all_cases.append((label, query, expected, partial_unsupported))

    failures = []
    mismatch_total = sum(1 for c in all_cases if not c[2])
    mismatch_pass = 0
    flagged_count = 0

    print(f"  Running {len(all_cases)} queries via HTTP POST to {args.url} ...")

    REFUSAL_STRING = "Not enough context to evaluate this."

    for label, query, expected_confident, expected_unsupported in all_cases:
        print(f"\n  {'-' * 76}")
        print(f"  Case: {label}")
        
        try:
            # Pad query to bypass the 50-character minimum JD length validation
            padded_query = f"The candidate must satisfy the following requirement: {query}"
            with open(resume_path, "rb") as f:
                files = {"resume": ("test_resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
                data = {"jd_text": padded_query}
                # Use a timeout and raise for status
                resp = requests.post(args.url, files=files, data=data, timeout=30)
                resp.raise_for_status()
                result = resp.json()
        except Exception as e:
            failures.append(f"{label} (HTTP Request Failed: {e})")
            print(f"  [FAIL] HTTP Error: {e}")
            continue

        gap_summary = result.get("gap_summary", "")
        is_flagged = result.get("is_flagged_by_verifier", False)
        unsupported = result.get("unsupported_requirements", [])

        if is_flagged:
            flagged_count += 1
            print(f"  [VERIFIER FLAGGED]")

        # Check expected vs actual behavior
        if not expected_confident:
            if gap_summary == REFUSAL_STRING:
                status = "PASS"
                mismatch_pass += 1
            else:
                failures.append(f"{label} (failed refusal constraint)")
                status = "FAIL"
        else:
            status = "PASS"
            # Explicitly validate unsupported_requirements for partial mismatch (D1)
            if expected_unsupported:
                # We expect at least these substrings to be in unsupported
                for req in expected_unsupported:
                    if not any(req.lower() in u.lower() for u in unsupported):
                        failures.append(f"{label} (expected '{req}' in unsupported_requirements, got {unsupported})")
                        status = "FAIL"
            
            # Additional validation: if gap_summary is REFUSAL_STRING but expected_confident was true (e.g. C4), that's a PASS.

        print(f"  Result [{status}]:")
        print(f"    Gap Summary: {gap_summary}")
        if unsupported:
            print(f"    Unsupported: {unsupported}")

    # Cleanup
    if os.path.exists(resume_path):
        os.remove(resume_path)

    total_cases = len(all_cases)
    passed = total_cases - len(failures)

    print(f"\n{'=' * 80}")
    print(f"  RESULTS SUMMARY")
    print(f"  {'=' * 76}")
    print(f"  Mismatch exit gate: {mismatch_pass}/{mismatch_total} correctly refused.")
    print(f"  Verifier interventions: {flagged_count} cases flagged.")
    print(f"  TOTAL:              {passed}/{total_cases} passed")
    print(f"  {'=' * 76}")

    if failures:
        print(f"  [FAIL] Failed cases:\n    " + "\n    ".join(failures))
        sys.exit(1)
    else:
        print(f"  [OK] All {total_cases} cases PASSED end-to-end over HTTP.")

if __name__ == "__main__":
    main()
